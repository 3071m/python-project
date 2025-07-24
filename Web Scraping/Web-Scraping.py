import requests
from bs4 import BeautifulSoup
from docx import Document

def scrape_multiple_pages_to_doc(base_url, start_page, end_page, output_filename='products.docx'):
    doc = Document()
    doc.add_heading('CeramicDeeDee Products (Page 2–43)', level=1)

    # สร้างตารางในเอกสาร
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Type'
    hdr_cells[1].text = 'Text'
    hdr_cells[2].text = 'URL'

    for page in range(start_page, end_page + 1):
        url = f"{base_url}{page}/"
        print(f"กำลังดึงข้อมูลจากหน้า: {url}")

        try:
            response = requests.get(url, timeout=10)
            response.raise_for_status()
        except Exception as e:
            print(f"❌ ข้ามหน้า {page} เนื่องจากเกิดข้อผิดพลาด: {e}")
            continue

        soup = BeautifulSoup(response.text, 'html.parser')

        # ลิงก์
        for a_tag in soup.find_all('a', href=True):
            href = a_tag['href']
            text = a_tag.get_text(strip=True)
            if text or href:
                row = table.add_row().cells
                row[0].text = 'Link'
                row[1].text = text
                row[2].text = href

        # ข้อความจาก <p>
        for p_tag in soup.find_all('p'):
            text = p_tag.get_text(strip=True)
            if text:
                row = table.add_row().cells
                row[0].text = 'Text'
                row[1].text = text
                row[2].text = ''

    doc.save(output_filename)
    print(f"\n✅ บันทึกข้อมูลลงไฟล์ {output_filename} เรียบร้อยแล้ว")

if __name__ == '__main__':
    base_url = "https://ceramicdeedee.com/products/page/"
    scrape_multiple_pages_to_doc(base_url, start_page=2, end_page=43)
